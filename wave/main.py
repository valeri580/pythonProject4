#!/usr/bin/env python3
import os
import json
import shutil
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple

import typer
from rich.console import Console
from rich.table import Table
from rich.progress import track
from dotenv import load_dotenv
import httpx
import subprocess
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Инициализация ---
app = typer.Typer(help="Терминальный аудио-помощник: TTS, классификация (music/speech/noise), транскрибация, отчёты.")
console = Console()
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
HF_API_TOKEN   = os.getenv("HF_API_TOKEN", "")

# Внешние API
from openai import OpenAI
openai_client = OpenAI(api_key=OPENAI_API_KEY)

HF_YAMNET_URL = "https://api-inference.huggingface.co/models/qualcomm/YamNet"
HF_HEADERS    = {"Authorization": f"Bearer {HF_API_TOKEN}"} if HF_API_TOKEN else {}

# Директории проекта
DATA_DIR   = Path("data")
OUT_DIR    = Path("out")
REPORTS_DIR= Path("reports")
for d in [DATA_DIR, OUT_DIR, REPORTS_DIR]:
    d.mkdir(exist_ok=True)

RESULTS_JSON = OUT_DIR / "classification_results.json"


# --- Утилиты ---

def ensure_audio_wav_16k_mono(src_path: Path) -> Path:
    """
    Приводит аудио к WAV 16 kHz mono (если нужно). Возвращает путь к конвертированному файлу (временный в OUT_DIR).
    Ищет ffmpeg в системе или в локальной папке ffmpeg_portable.
    """
    out_path = OUT_DIR / (src_path.stem + "_16k_mono.wav")
    
    # Ищем FFmpeg в разных местах
    ffmpeg_paths = [
        "ffmpeg",  # в PATH
        "ffmpeg_portable/ffmpeg.exe",  # локальная папка
        "C:/ffmpeg/bin/ffmpeg.exe",  # стандартная установка
        "C:/ffmpeg/ffmpeg.exe"  # альтернативная установка
    ]
    
    ffmpeg_cmd = None
    for path in ffmpeg_paths:
        try:
            result = subprocess.run([path, "-version"], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                ffmpeg_cmd = path
                break
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if not ffmpeg_cmd:
        raise RuntimeError("FFmpeg не найден. Установите FFmpeg или запустите ffmpeg_download.bat")
    
    # Используем найденный ffmpeg
    cmd = [
        ffmpeg_cmd, "-y",  # -y для перезаписи файлов
        "-i", str(src_path),
        "-ar", "16000",  # sample rate 16kHz
        "-ac", "1",      # mono (1 channel)
        "-acodec", "pcm_s16le",  # 16-bit PCM
        str(out_path)
    ]
    
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return out_path
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Ошибка конвертации аудио с ffmpeg: {e.stderr}")
    except FileNotFoundError:
        raise RuntimeError(f"FFmpeg не найден по пути: {ffmpeg_cmd}")


# Наборы ключевых слов для свёртки YAMNet-меток в 3 класса.
MUSIC_KEYS  = {
    "music", "musical", "instrument", "sing", "choir", "song",
    "guitar","piano","violin","drum","bass","flute","saxophone","trumpet","cello","harmonica","organ","synth",
    "orchestra","hip hop","rock","pop","techno","edm","metal","reggae","blues","jazz","classical"
}
SPEECH_KEYS = {
    "speech","conversation","narration","monologue","dialog","dialogue",
    "telephone","talking","babbling","whisper","shout","speech synthesizer","broadcast"
}

def fold_labels_to_triple(labels: List[Dict[str, Any]]) -> Tuple[str, Dict[str, float]]:
    """
    Преобразует список {label, score} от YAMNet в итоговую категорию: music / speech / noise.
    Возвращает (final_label, scores_dict).
    """
    scores = {"music": 0.0, "speech": 0.0, "noise": 0.0}
    for item in labels:
        lab = item.get("label", "").lower()
        sc  = float(item.get("score", 0.0))
        if any(k in lab for k in MUSIC_KEYS):
            scores["music"] += sc
        elif any(k in lab for k in SPEECH_KEYS):
            scores["speech"] += sc
        else:
            scores["noise"] += sc
    final = max(scores.items(), key=lambda kv: kv[1])[0]
    return final, scores


def yamnet_classify(file_path: Path, top_k: int = 10) -> List[Dict[str, Any]]:
    """
    Отправляет аудио в HF Inference API (YAMNet) и возвращает top-k меток c вероятностями.
    """
    if not HF_API_TOKEN:
        raise RuntimeError("Не задан HF_API_TOKEN в .env — требуется для вызова Inference API.")

    with open(file_path, "rb") as f:
        data = f.read()

    with httpx.Client(timeout=60.0) as client:
        r = client.post(HF_YAMNET_URL, headers=HF_HEADERS, content=data)
        r.raise_for_status()
        resp = r.json()

    # Ответ Inference API обычно: [{label: "...", score: float}, ...]
    # На всякий случай упорядочим и обрежем top-k:
    if isinstance(resp, list):
        resp_sorted = sorted(resp, key=lambda d: d.get("score", 0.0), reverse=True)
        return resp_sorted[:top_k]
    # Иногда API возвращает dict с предупреждением (например, cold start) — повторим:
    if isinstance(resp, dict) and "error" in resp:
        raise RuntimeError(f"Hugging Face Inference API error: {resp.get('error')}")
    return []


def transcribe_openai(file_path: Path, prompt: Optional[str] = None) -> Dict[str, Any]:
    """
    Транскрибация аудио через OpenAI Audio API (ASR).
    Поддерживает MP3, WAV, M4A, FLAC и другие форматы напрямую.
    """
    if not OPENAI_API_KEY:
        raise RuntimeError("Не задан OPENAI_API_KEY в .env")

    # OpenAI Whisper поддерживает множество форматов напрямую, конвертация не всегда нужна
    try:
        with open(file_path, "rb") as f:
            resp = openai_client.audio.transcriptions.create(
                model="whisper-1",  # Стандартная модель Whisper
                file=f,
                prompt=prompt or "",
                response_format="verbose_json"  # получим тайминги, если доступны
            )
        # SDK возвращает объект со свойствами; преобразуем к dict
        return json.loads(resp.model_dump_json())
    except Exception as e:
        # Если прямая отправка не сработала, пробуем конвертировать
        if "FFmpeg не найден" not in str(e):
            try:
                wav = ensure_audio_wav_16k_mono(file_path)
                with open(wav, "rb") as f:
                    resp = openai_client.audio.transcriptions.create(
                        model="whisper-1",
                        file=f,
                        prompt=prompt or "",
                        response_format="verbose_json"
                    )
                return json.loads(resp.model_dump_json())
            except Exception:
                pass
        raise RuntimeError(f"Ошибка транскрибации: {e}")


def tts_openai(text: str, voice: str, out_path: Path, fmt: str = "mp3"):
    """
    Озвучка текста через OpenAI Audio API (TTS).
    Модель: gpt-4o-mini-tts
    """
    if not OPENAI_API_KEY:
        raise RuntimeError("Не задан OPENAI_API_KEY в .env")

    # Используем стандартный API OpenAI TTS
    try:
        response = openai_client.audio.speech.create(
            model="tts-1",  # Используем стандартную модель TTS
            voice=voice,
            input=text,
            response_format=fmt
        )
        
        # Сохраняем аудио в файл
        with open(out_path, "wb") as f:
            f.write(response.content)
            
    except Exception as e:
        raise RuntimeError(f"Ошибка при генерации TTS: {e}")


def save_results(results: List[Dict[str, Any]]):
    OUT_DIR.mkdir(exist_ok=True)
    with open(RESULTS_JSON, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)


def load_results() -> List[Dict[str, Any]]:
    if RESULTS_JSON.exists():
        return json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    return []


def show_table(results: List[Dict[str, Any]]):
    table = Table(title="Результаты классификации аудио")
    table.add_column("Файл", style="cyan")
    table.add_column("Класс", style="magenta")
    table.add_column("Уверенность", justify="right")
    table.add_column("Top-метки (YAMNet)", overflow="fold")

    for r in results:
        top_labels_str = ", ".join([f"{l['label']} ({l['score']:.2f})" for l in r.get("top_labels", [])[:5]])
        conf = r.get("scores", {}).get(r.get("final_label",""), 0.0)
        table.add_row(r["file"], r["final_label"], f"{conf:.2f}", top_labels_str)
    console.print(table)


def export_html(results: List[Dict[str, Any]], outfile: Path):
    rows = []
    for r in results:
        s = r.get("scores", {})
        rows.append(f"""
        <tr>
          <td>{r['file']}</td>
          <td class="badge {r['final_label']}">{r['final_label']}</td>
          <td>
            <div class="bar">
              <span style="width:{s.get('music',0)*100:.1f}%"></span>
            </div>
            <small>music: {s.get('music',0):.2f}</small>
          </td>
          <td>
            <div class="bar">
              <span style="width:{s.get('speech',0)*100:.1f}%"></span>
            </div>
            <small>speech: {s.get('speech',0):.2f}</small>
          </td>
          <td>
            <div class="bar">
              <span style="width:{s.get('noise',0)*100:.1f}%"></span>
            </div>
            <small>noise: {s.get('noise',0):.2f}</small>
          </td>
        </tr>
        """)

    html = f"""<!doctype html>
<html lang="ru">
<head>
<meta charset="utf-8"/>
<title>Отчёт по классификации аудио</title>
<style>
 body {{ font-family: system-ui, -apple-system, Segoe UI, Roboto, sans-serif; margin: 24px; }}
 h1 {{ margin-bottom: 8px; }}
 table {{ border-collapse: collapse; width: 100%; }}
 th, td {{ border-bottom: 1px solid #eee; text-align: left; padding: 8px; vertical-align: top; }}
 .badge {{ display:inline-block; padding: 2px 8px; border-radius: 999px; color:#fff; font-size: 12px; }}
 .badge.music  {{ background: #5b8; }}
 .badge.speech {{ background: #58b; }}
 .badge.noise  {{ background: #b85; }}
 .bar {{ background:#f1f1f1; height:10px; border-radius:6px; overflow:hidden; }}
 .bar span {{ display:block; height:10px; background:#888; }}
 small {{ color:#666; }}
</style>
</head>
<body>
<h1>Отчёт по классификации аудио</h1>
<p>Сводка по файлам: музыка / речь / шум (по данным YAMNet, агрегированным в 3 класса).</p>
<table>
  <thead>
    <tr>
      <th>Файл</th><th>Класс</th>
      <th>music</th><th>speech</th><th>noise</th>
    </tr>
  </thead>
  <tbody>
    {''.join(rows)}
  </tbody>
</table>
</body>
</html>
"""
    outfile.write_text(html, encoding="utf-8")
    console.print(f"[green]HTML-отчёт сохранён:[/green] {outfile}")


# --- Команды CLI ---

@app.command()
def ingest(paths: List[Path] = typer.Argument(..., help="Пути к файлам или папкам с аудио"),
           recurse: bool = typer.Option(True, help="Рекурсивно искать аудио в папках"),
           exts: str = typer.Option("wav,mp3,m4a,flac,ogg", help="Расширения через запятую")):
    """
    "Загрузка" (ингест) аудио в локальную папку проекта ./data/
    """
    exts_set = {e.strip().lower() for e in exts.split(",")}
    copied = []
    for p in paths:
        if p.is_dir():
            it = p.rglob("*") if recurse else p.iterdir()
            for f in it:
                if f.is_file() and f.suffix.lower().lstrip(".") in exts_set:
                    dest = DATA_DIR / f.name
                    shutil.copy2(f, dest)
                    copied.append(dest.name)
        elif p.is_file() and p.suffix.lower().lstrip(".") in exts_set:
            dest = DATA_DIR / p.name
            shutil.copy2(p, dest)
            copied.append(dest.name)
    console.print(f"[green]Скопировано:[/green] {len(copied)} файлов → ./data")
    for name in copied:
        console.print(f"  • {name}")


@app.command()
def tts(
    text: str = typer.Argument(..., help="Текст для озвучки"),
    voice: str = typer.Option("alloy", help="Голос TTS (например: alloy, verse, coral и т.п.)"),
    out: Path = typer.Option(OUT_DIR / "tts_output.mp3", help="Файл для сохранения"),
    fmt: str = typer.Option("mp3", help="Формат (mp3|wav)")
):
    """
    Озвучивает текст через OpenAI TTS и сохраняет аудио.
    """
    tts_openai(text, voice, out, fmt)
    console.print(f"[green]Готово:[/green] {out}")


def split_large_audio(file_path: Path, max_size_mb: int = 20, segment_duration: int = 600) -> List[Path]:
    """
    Разбивает большой аудиофайл на части для обработки в OpenAI API.
    max_size_mb: максимальный размер части в МБ
    segment_duration: длительность сегмента в секундах (по умолчанию 10 минут)
    """
    file_size_mb = file_path.stat().st_size / (1024 * 1024)
    
    if file_size_mb <= max_size_mb:
        return [file_path]  # Файл не нужно разбивать
    
    console.print(f"[yellow]Файл большой ({file_size_mb:.1f} МБ), разбиваем на части...[/yellow]")
    
    # Ищем FFmpeg
    ffmpeg_paths = [
        "ffmpeg",
        "ffmpeg_portable/ffmpeg.exe", 
        "C:/ffmpeg/ffmpeg.exe",
        "C:/ffmpeg/bin/ffmpeg.exe"
    ]
    
    ffmpeg_cmd = None
    for path in ffmpeg_paths:
        try:
            result = subprocess.run([path, "-version"], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                ffmpeg_cmd = path
                break
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if not ffmpeg_cmd:
        raise RuntimeError("FFmpeg не найден. Требуется для разбивки больших файлов.")
    
    # Создаем папку для частей
    parts_dir = OUT_DIR / f"{file_path.stem}_parts"
    parts_dir.mkdir(exist_ok=True)
    
    # Разбиваем файл на части
    output_pattern = parts_dir / f"{file_path.stem}_part_%03d.mp3"
    cmd = [
        ffmpeg_cmd, "-i", str(file_path),
        "-f", "segment", "-segment_time", str(segment_duration),
        "-c", "copy", str(output_pattern)
    ]
    
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        
        # Находим созданные части
        parts = sorted(parts_dir.glob(f"{file_path.stem}_part_*.mp3"))
        console.print(f"[green]Создано частей:[/green] {len(parts)}")
        
        return parts
        
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Ошибка разбивки файла: {e.stderr}")


@app.command()
def transcribe(
    file: Path = typer.Argument(..., help="Аудиофайл для транскрибации"),
    prompt: Optional[str] = typer.Option(None, help="Контекст/словарь имён для ASR"),
    out: Optional[Path] = typer.Option(None, help="Куда сохранить текст (если не задать — выведем в консоль)"),
    split: bool = typer.Option(True, help="Автоматически разбивать большие файлы"),
    word: bool = typer.Option(False, help="Дополнительно создать Word документ")
):
    """
    Транскрибация аудио с помощью OpenAI (ASR). 
    Автоматически разбивает большие файлы на части.
    """
    if split:
        # Разбиваем файл на части если нужно
        parts = split_large_audio(file)
        
        if len(parts) == 1:
            # Файл маленький, обрабатываем как обычно
            res = transcribe_openai(file, prompt=prompt)
            text = (res.get("text") or "").strip()
        else:
            # Обрабатываем части
            all_texts = []
            for i, part in enumerate(track(parts, description="Транскрибация частей...")):
                try:
                    res = transcribe_openai(part, prompt=prompt)
                    part_text = (res.get("text") or "").strip()
                    if part_text:
                        all_texts.append(f"=== Часть {i+1} ===\n{part_text}")
                        console.print(f"[green]Часть {i+1}/{len(parts)} готова[/green]")
                except Exception as e:
                    console.print(f"[red]Ошибка в части {i+1}:[/red] {e}")
                    all_texts.append(f"=== Часть {i+1} ===\n[ОШИБКА: {e}]")
            
            text = "\n\n".join(all_texts)
            
            # Очищаем временные файлы
            for part in parts[1:]:  # Оставляем оригинальный файл
                try:
                    part.unlink()
                except:
                    pass
            
            # Удаляем папку с частями если она пуста
            parts_dir = parts[0].parent if len(parts) > 1 else None
            if parts_dir and parts_dir != file.parent:
                try:
                    parts_dir.rmdir()
                except:
                    pass
    else:
        # Обычная транскрибация без разбивки
        res = transcribe_openai(file, prompt=prompt)
        text = (res.get("text") or "").strip()
    
    if out:
        out.write_text(text, encoding="utf-8")
        console.print(f"[green]Транскрипт сохранён:[/green] {out}")
        
        # Создаем Word документ если запрошено
        if word:
            word_file = out.with_suffix('.docx')
            title = f"Транскрипт: {file.stem}"
            create_word_document(out, word_file, title)
            console.print(f"[green]Word документ создан:[/green] {word_file}")
    else:
        console.rule("[bold]Транскрипт[/bold]")
        console.print(text or "[пусто]")


@app.command()
def classify(
    inputs: List[Path] = typer.Argument(..., help="Файлы или папки"),
    recurse: bool = typer.Option(True, help="Рекурсивный поиск по папкам"),
    save: bool = typer.Option(True, help="Сохранить summary в JSON")
):
    """
    Классификация аудио через YAMNet (HF Inference API) ⇒ music / speech / noise.
    """
    files: List[Path] = []
    for p in inputs:
        if p.is_dir():
            it = p.rglob("*") if recurse else p.iterdir()
            for f in it:
                if f.is_file():
                    files.append(f)
        elif p.is_file():
            files.append(p)

    results: List[Dict[str, Any]] = []
    for f in track(files, description="Классификация..."):
        try:
            # YAMNet принимает разные форматы; для надёжности приводим к WAV 16k mono
            wav = ensure_audio_wav_16k_mono(f)
            labels = yamnet_classify(wav, top_k=12)
            final_label, scores = fold_labels_to_triple(labels)
            results.append({
                "file": f.name,
                "final_label": final_label,
                "scores": scores,
                "top_labels": labels
            })
        except Exception as e:
            console.print(f"[red]Ошибка:[/red] {f} — {e}")

    if save:
        save_results(results)
    show_table(results)


def create_word_document(text_file: Path, output_file: Path, title: str = "Транскрипт аудиозаписи"):
    """
    Создает Word документ из текстового файла с красивым форматированием.
    """
    # Читаем исходный текст
    content = text_file.read_text(encoding="utf-8")
    
    # Создаем новый документ
    doc = Document()
    
    # Добавляем заголовок
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем информацию о файле
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run(f"Исходный файл: {text_file.name}\n").bold = True
    info_paragraph.add_run(f"Дата создания: {text_file.stat().st_mtime}\n")
    info_paragraph.add_run(f"Размер транскрипта: {len(content)} символов\n")
    
    # Добавляем разделитель
    doc.add_paragraph("─" * 50)
    
    # Обрабатываем содержимое
    parts = content.split("=== Часть")
    
    if len(parts) > 1:
        # Файл разбит на части
        doc.add_paragraph("Документ содержит транскрипт аудиозаписи, разбитой на части:")
        
        for i, part in enumerate(parts[1:], 1):  # Пропускаем первую пустую часть
            lines = part.strip().split('\n')
            if not lines:
                continue
                
            # Заголовок части
            part_title = f"Часть {i}"
            if lines[0].strip().endswith("==="):
                part_title = lines[0].replace("===", "").strip()
                lines = lines[1:]  # Убираем строку с заголовком
            
            doc.add_heading(part_title, level=1)
            
            # Текст части
            part_text = '\n'.join(lines).strip()
            if part_text:
                # Разбиваем на абзацы для лучшей читаемости
                paragraphs = part_text.split('\n\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        p = doc.add_paragraph(paragraph_text.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Добавляем небольшой отступ между частями
            doc.add_paragraph("")
    else:
        # Обычный файл без разбивки
        paragraphs = content.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Сохраняем документ
    doc.save(output_file)
    return output_file


@app.command()
def to_word(
    text_file: Path = typer.Argument(..., help="Текстовый файл для конвертации"),
    output: Optional[Path] = typer.Option(None, help="Имя выходного Word файла"),
    title: str = typer.Option("Транскрипт аудиозаписи", help="Заголовок документа")
):
    """
    Конвертирует текстовый транскрипт в красиво оформленный Word документ.
    """
    if not text_file.exists():
        console.print(f"[red]Файл не найден:[/red] {text_file}")
        raise typer.Exit(code=1)
    
    if not output:
        output = text_file.with_suffix('.docx')
    
    try:
        result_file = create_word_document(text_file, output, title)
        console.print(f"[green]Word документ создан:[/green] {result_file}")
        console.print(f"[cyan]Размер файла:[/cyan] {result_file.stat().st_size / 1024:.1f} КБ")
    except Exception as e:
        console.print(f"[red]Ошибка создания документа:[/red] {e}")
        raise typer.Exit(code=1)


@app.command()
def report(
    html: Optional[Path] = typer.Option(REPORTS_DIR / "classification_report.html", help="Путь к HTML-отчёту")
):
    """
    Печать результатов и (опционально) экспорт HTML-отчёта.
    """
    results = load_results()
    if not results:
        console.print("[yellow]Нет сохранённых результатов. Сначала запустите 'classify'.[/yellow]")
        raise typer.Exit(code=0)
    show_table(results)
    if html:
        export_html(results, html)


if __name__ == "__main__":
    app()
